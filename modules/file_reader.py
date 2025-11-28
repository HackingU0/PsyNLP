"""
File Reader Module
Supports reading text from .md, .txt, and .docx files
"""

import os
from pathlib import Path
from io import BytesIO


def read_markdown(file_path) -> str:
    """Read markdown file - supports both path string and file-like object"""
    if isinstance(file_path, str):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        # Handle file-like object (e.g., Streamlit UploadedFile)
        return file_path.read().decode("utf-8")


def read_text(file_path) -> str:
    """Read plain text file - supports both path string and file-like object"""
    if isinstance(file_path, str):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        # Handle file-like object
        return file_path.read().decode("utf-8")


def read_docx(file_path) -> str:
    """Read DOCX file - supports both path string and file-like object"""
    try:
        from docx import Document

        if isinstance(file_path, str):
            doc = Document(file_path)
        else:
            doc = Document(BytesIO(file_path.read()))

        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n\n".join(text)
    except ImportError:
        raise ImportError(
            "python-docx is required to read .docx files. "
            "Install it with: pip install python-docx"
        )


def read_file(file_path) -> str:
    """
    Read file based on extension
    Supports: .md, .txt, .docx
    Works with both file paths (str) and file-like objects (UploadedFile)
    """
    # Get extension
    if isinstance(file_path, str):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        ext = Path(file_path).suffix.lower()
    else:
        # Handle file-like object (get name from object)
        ext = Path(file_path.name).suffix.lower()

    if ext == ".md":
        return read_markdown(file_path)
    elif ext == ".txt":
        return read_text(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}\n" "Supported formats: .md, .txt, .docx"
        )


def get_file_info(file_path) -> dict:
    """Get file information"""
    stat = os.stat(file_path)
    return {
        "name": os.path.basename(file_path),
        "path": file_path,
        "size": stat.st_size,
        "extension": Path(file_path).suffix.lower(),
    }


# Test
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            content = read_file(file_path)
            info = get_file_info(file_path)

            print(f"File: {info['name']}")
            print(f"Size: {info['size']} bytes")
            print(f"Extension: {info['extension']}")
            print(f"\nContent preview:")
            print(content[:200] + "..." if len(content) > 200 else content)
        except Exception as e:
            print(f"Error: {e}")
    else:
        print("Usage: python file_reader.py <file_path>")
